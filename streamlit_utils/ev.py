import os
import json
import uuid
import time
from io import BytesIO
from dotenv import load_dotenv
import json
import gpt_utils, db_utils, common_utils, pinecone_utils, filemapping_utils, pdf_utils
from api_utils import get_current_evaluator_prompts
from db_utils import log_analyzer_custom_evaluator_data
from s3_bucket_utils import upload_fileobj_to_s3
import s3_bucket_utils
import docx
from pdfminer.high_level import extract_text
import concurrent.futures
import traceback
from logger import streamlit_logger as logger


load_dotenv(override=True)

text_extraction = os.getenv("text_extraction")

possible_models = json.loads(os.getenv("possible_models"))

prompt_labels = ["P1", "P2", "P3", "P4", "P5"]

pdf_mappings = filemapping_utils.get_pdf_mappings()


def extract_text_from_pdf(uploaded_file, st):
    try:
        with BytesIO() as buffer:
            buffer.write(uploaded_file.read())
            buffer.seek(0)
            return extract_text(buffer)
    except Exception as e:
        st.error(f"Failed to extract text from the PDF: {e}")
        return None


def evaluator_prompts_panel(st):
    st.write("Upload Proposal file:")
    proposal_file = st.file_uploader("Choose a file", type=["pdf", "txt", "docx"], key="proposal")

    st.write("Upload TOR file:")
    tor_file = st.file_uploader("Choose a file", type=["pdf", "txt", "docx"], key="tor")

    proposal_content = ""
    proposal_pdf_name = ""

    tor_content = ""
    tor_pdf_name = ""

    if proposal_file:
        proposal_content = proposal_file.getvalue()
        proposal_pdf_name = proposal_file.name

    if tor_file:
        tor_content = tor_file.getvalue()
        tor_pdf_name = tor_file.name
        
    nature_options = [
        "Policy Document",
        "Investment or grant proposal",
        "Research draft or proposal",
        "Program design Document",
        "Strategy recommendations",
        "Media article or draft",
        "School or college course outline",
        "MEL approach",
        "Product or service design"
    ]
    
    nature_default = "Program design Document"

    nature = st.selectbox("Nature of Document", nature_options, index=nature_options.index(nature_default))

    organization_id = st.text_input(label="Enter Organization ID. Eg. UNICEF, BMGF")
    
    org_guideline_id = st.text_input(label="Enter Org Guideline ID.")
    org_guideline_id = org_guideline_id if org_guideline_id else None
    
    prompt_fields = {
        'P_Internal': ['P_Internal.F1', 'P_Internal.F2']
    }

    selected_prompts_labels = st.multiselect("Prompts", ["P_Internal",])
    
    available_flows = []
    selected_flows = []

    for prompt in selected_prompts_labels:
        available_flows.extend(prompt_fields[prompt])

    if available_flows:
        selected_flows = st.multiselect('Select Fields', available_flows)

    #selected_model_key = st.selectbox('Prompt Model', ('gpt-4','gpt-3.5'))
    model_options = {'gpt-4-omni': 'gpt-4', 'gpt-3.5': 'gpt-3.5'}
    selected_display_name = st.selectbox('Prompt Model', list(model_options.keys()))
    selected_model_key = model_options[selected_display_name]

    selected_model = possible_models[selected_model_key]

    submit = st.button("Generate Prompts")

    return proposal_content, proposal_pdf_name, tor_content, tor_pdf_name, nature, organization_id, org_guideline_id, selected_prompts_labels, selected_flows, selected_model, submit


def generate_evaluator_prompts(st, proposal_content, proposal_pdf_name, tor_content, tor_pdf_name, nature, organization_id, org_guideline_id, selected_prompt_labels, selected_flows, selected_model):
    try:
        proposal_summary_text = ""
        proposal_text = ""
        tor_text = None
        tor_summary_text = None
        prompts_associated = {}
        input_tokens={}
        output_tokens={}
        
        with st.spinner("Fetching Evaluator prompts..."):
            start_time = time.time()
            #evaluator_prompts, evaluator_wisdom, evaluator_section_titles, evaluator_dependency_graph = db_utils.get_evaluator_prompts_multithreaded(partition_type=selected_prompt_labels, doc_type=nature, organization_id=organization_id, org_guideline_id=org_guideline_id, st=st)
            print(f"Selected Prompt Labels: {selected_prompt_labels}")
            evaluator_prompts, evaluator_prompt_flows = get_current_evaluator_prompts(selected_prompt_labels, nature, organization_id, org_guideline_id)
            if not evaluator_prompts:
                st.error("No prompts found for the organization: " + organization_id)
                logger.error("No prompts found for the organization: " + organization_id)
                return -1
            logger.info(f"Fetched evaluator prompts in {time.time() - start_time:.2f} seconds.")

        logger.info("Processing Proposal and TOR files concurrently...")
        
        def process_proposal(proposal_content, proposal_pdf_name, text_extraction, selected_model, nature):
            proposal_text = ""
            if proposal_content:
                if proposal_pdf_name.endswith(".docx"):
                    doc = docx.Document(BytesIO(proposal_content))
                    proposal_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                elif proposal_pdf_name.endswith(".pdf"):
                    if text_extraction == "PDF Reader":
                        proposal_text = pdf_utils.extract_text_from_pdf(BytesIO(proposal_content))
                    else:
                        proposal_text = pdf_utils.extract_text_llama_parse(proposal_content, proposal_pdf_name)
                elif proposal_pdf_name.endswith(".txt"):
                    proposal_text = proposal_content.decode("utf-8")

                s3_par_url = upload_fileobj_to_s3(BytesIO(proposal_content), proposal_pdf_name)
                
                proposal_summary_text, proposal_summary_prompt = gpt_utils.get_summary_streamlit(proposal_text, selected_model, nature)
                
                return s3_par_url, proposal_text, proposal_summary_text, proposal_summary_prompt
            return "", ""

        def process_tor(tor_content, tor_pdf_name, text_extraction, selected_model, nature, organization_id):
            tor_text = ""
            if tor_content:
                if tor_pdf_name.endswith(".docx"):
                    doc = docx.Document(BytesIO(tor_content))
                    tor_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                elif tor_pdf_name.endswith(".pdf"):
                    if text_extraction == "PDF Reader":
                        tor_text = pdf_utils.extract_text_from_pdf(BytesIO(tor_content))
                    else:
                        tor_text = pdf_utils.extract_text_llama_parse(tor_content, tor_pdf_name)
                elif tor_pdf_name.endswith(".txt"):
                    tor_text = tor_content.decode("utf-8")
                
                tor_summary_text, tor_summary_prompt = gpt_utils.get_tor_summary_streamlit(tor_text, selected_model, nature, organization_id)
                return tor_summary_text, tor_summary_prompt
            return "", ""

        with st.spinner("Processing Input Files..."):
            start_time = time.time()
            
            with concurrent.futures.ThreadPoolExecutor() as executor:
                proposal_future = executor.submit(process_proposal, proposal_content, proposal_pdf_name, text_extraction, selected_model, nature)
                tor_future = executor.submit(process_tor, tor_content, tor_pdf_name, text_extraction, selected_model, nature, organization_id)
                
                s3_par_url, proposal_summary_text, proposal_summary_prompt = proposal_future.result()
                tor_summary_text, tor_summary_prompt = tor_future.result()

                st.session_state['proposal_pdf_url'] = s3_par_url if s3_par_url else ""

        logger.info(f"Processed Proposal and TOR files in {time.time() - start_time:.2f} seconds.")
            
        with st.spinner("Generating Evaluator Prompts..."):
            evaluator_dependencies = {}
            
            for key, value in evaluator_prompts.items():
                
                # Associate prompts with prompt labels
                prompts_associated[key] = {
                    "base_prompt": value["base_prompt"],
                    "customization_prompt": value["customization_prompt"]
                }
                if value["additional_dependencies"]:
                    evaluator_dependencies[key] = value["additional_dependencies"]
                else:
                    evaluator_dependencies[key] = []
                    
            generated_evaluator_prompts, _, _ = gpt_utils.generate_evaluator_prompts(selected_model, nature, "", tor_summary_text, evaluator_prompts)


            evaluator_flows = {}
            for key in evaluator_prompt_flows.keys():
                dictionary = evaluator_prompt_flows[key]
                evaluator_flows[f"{key}"] = {}
                for sub_key in dictionary.keys():
                    if sub_key in selected_flows:
                        evaluator_flows[f"{key}"][sub_key] = evaluator_prompt_flows[f"{key}"][sub_key]

        # if selected_prompt_labels:
        #     if "P_Internal" in selected_prompt_labels:
        #         generated_evaluator_prompts["P_Internal"]["generated_prompt"]=evaluator_prompts["P_Internal"]["base_prompt"]
        # else:
        #     generated_evaluator_prompts["P_Internal"]["generated_prompt"]=evaluator_prompts["P_Internal"]["base_prompt"]

        if "proposal_pdf_url" not in st.session_state:
            st.session_state["proposal_pdf_url"] = s3_par_url
            
        st.session_state["input_tokens"] = input_tokens
        st.session_state["output_tokens"] = output_tokens
        st.session_state['evaluator_prompts'] = prompts_associated
        st.session_state['evaluator_flows'] = evaluator_flows
        st.session_state["evaluator_dependencies"] = evaluator_dependencies
        st.session_state['generated_evaluator_prompts'] = generated_evaluator_prompts
        st.session_state['proposal_summary_text'] = proposal_summary_text
        st.session_state['proposal_text'] = proposal_text
        st.session_state['tor_summary_text'] = tor_summary_text
        st.session_state['tor_text'] = tor_text
        st.session_state['proposal_pdf_name'] = proposal_pdf_name
        st.session_state['tor_pdf_name'] = tor_pdf_name
            # start_time = time.time()
            # for key, value in evaluator_prompts.items():
            #     prompts_associated[key] = {
            #         "base_prompt": value["base_prompt"],
            #         "customization_prompt": value["customization_prompt"]
            #     }
            # generated_evaluator_prompts = common_utils.generate_evaluator_prompts(selected_model, nature, tor_summary_text, evaluator_prompts)
            
            # if selected_prompt_labels:
            #     if "P_Internal" in selected_prompt_labels:
            #         generated_evaluator_prompts["P_Internal"] = evaluator_prompts["P_Internal"]["base_prompt"]
            # else:
            #     generated_evaluator_prompts["P_Internal"] = evaluator_prompts["P_Internal"]["base_prompt"]
            
            # logger.info(f"Generated evaluator prompts in {time.time() - start_time:.2f} seconds.")

        # st.session_state['evaluator_prompts'] = prompts_associated
        # st.session_state['generated_evaluator_prompts'] = generated_evaluator_prompts
        # st.session_state['proposal_summary_text'] = proposal_summary_text
        # st.session_state['proposal_text'] = proposal_text
        # st.session_state['tor_summary_text'] = tor_summary_text
        # st.session_state['tor_text'] = tor_text
        # st.session_state['proposal_pdf_name'] = proposal_pdf_name
        # st.session_state['tor_pdf_name'] = tor_pdf_name
        # st.session_state['evaluator_wisdom'] = evaluator_wisdom
        # st.session_state['evaluator_section_titles'] = evaluator_section_titles
        # st.session_state['evaluator_dependency_graph'] = evaluator_dependency_graph

    except Exception as e:
        logger.error(f"Error in generate_evaluator_prompts: {e}")
        traceback.print_exc()
        st.error(f"An error occurred in generate_evaluator_prompts: {e}")


def generate_evaluator_comments_old(st, nature, organization_id):
    try:
        start_time = time.time()
        
        logger.info("Loading session state variables...")
        proposal_pdf_name = st.session_state['proposal_pdf_name']
        tor_pdf_name = st.session_state['tor_pdf_name']
        used_analyze_prompts = st.session_state['analyze_prompts']
        generated_analyze_prompts = st.session_state['generated_analyze_prompts']
        wisdom = st.session_state['wisdom']
        used_evaluator_prompts = st.session_state['evaluator_prompts']
        generated_evaluator_prompts = st.session_state['generated_evaluator_prompts']
        proposal_pdf_url = st.session_state['proposal_pdf_url']
        proposal_summary_text = st.session_state['proposal_summary_text']
        proposal_text = st.session_state['proposal_text']
        tor_summary_text = st.session_state['tor_summary_text']
        tor_text = st.session_state['tor_text']
        wisdom = st.session_state['wisdom']
        evaluator_wisdom = st.session_state['evaluator_wisdom']
        prompts_section_title = st.session_state['prompts_section_title']
        evaluator_section_titles = st.session_state['evaluator_section_titles']
        dependency_graph = st.session_state['dependency_graph']
        evaluator_dependency_graph = st.session_state['evaluator_dependency_graph']
        logger.info("Loaded session state variables successfully.")
        
        logger.info("Displaying extracted data...")
        with st.expander("Proposal S3 Url"):
            st.markdown(f'<a href="{proposal_pdf_url}" target="_blank">{proposal_pdf_url}</a>', unsafe_allow_html=True)
        if used_analyze_prompts:
            with st.expander("Used Analyzer Prompts"):
                st.json(used_analyze_prompts)
        with st.expander("Used Evaluator Prompts"):
            st.json(used_evaluator_prompts)
        with st.expander("Proposal Summary"):
            st.text(proposal_summary_text)
        with st.expander("TOR Summary"):
            st.text(tor_summary_text)
        if generated_analyze_prompts:
            with st.expander("Generated Analyze Prompts"):
                st.json(generated_analyze_prompts)
        if dependency_graph:
            with st.expander("Dependency Graph"):
                st.json(dependency_graph)
        with st.expander("Generated Evaluator Prompts"):
            st.json(generated_evaluator_prompts)
        logger.info("Displayed extracted data successfully.")
        
        model_options = {'gpt-4-omni': 'gpt-4', 'gpt-3.5': 'gpt-3.5'}
        selected_display_name = st.selectbox('Evaluate Model', list(model_options.keys()))
        model_key = model_options[selected_display_name]

        model = possible_models[model_key]
        evaluate_submit = st.button("Evaluate")
            
        if evaluate_submit:
            logger.info("Starting evaluation comments generation...")
            evaluator_tokens_counter = common_utils.get_evaluator_tokens_counter()
            if generated_analyze_prompts:
                logger.info("Generating Analyze Comments...")
                with st.spinner("Generating Analyze Comments..."):
                    pinecone_analyze_filters = pinecone_utils.get_pinecone_analyzer_filters()
                    generated_analyze_comments, analyze_context_used = common_utils.generate_analyze_comments_evaluator(proposal_summary_text, pinecone_analyze_filters, pdf_mappings, model, generated_analyze_prompts, wisdom, nature, tokens_counter=evaluator_tokens_counter, dependency_graph=dependency_graph)
                    analyze_comments = common_utils.combine_comment(generated_analyze_comments)
                    logger.info("Analyzer comments generated successfully")
                    with st.expander("Generated Analyze Comments"):
                        st.text(analyze_comments)
                    with st.expander("Analyze Context Used"):
                        st.json(analyze_context_used)
            else:
                logger.info("No generated analyze prompts, skipping analysis comments generation")
                analyze_context_used = ""
                generated_analyze_comments = ""
                analyze_comments = ""
            
            
            with st.spinner("Generating Evaluate Comments..."):
                logger.info("Generating Evaluate Comments...")
                generated_evaluator_comments, p_internal_context, p_external_context, p_delta_context = gpt_utils.generate_evaluator_comments_concurrently_test(proposal_text, model, generated_evaluator_prompts, analyze_comments, evaluator_wisdom, evaluator_dependency_graph, tokens_counter=evaluator_tokens_counter)
                logger.info("Evaluation comments generated successfully")
                
                if p_internal_context:
                    with st.expander("P_Internal Context"):
                        st.json(p_internal_context)
                    with st.expander("P_Internal Section Title"):
                        st.text(evaluator_section_titles["P_Internal"])
                    with st.expander("Generated P_Internal Comments"):
                        st.text(generated_evaluator_comments["P_Internal"])
                
                if p_external_context:
                    with st.expander("P_External Context"):
                        st.json(p_external_context)
                    with st.expander("P_External Section Title"):
                        st.text(evaluator_section_titles["P_External"])
                    with st.expander("Generated P_External Comments"):
                        st.text(generated_evaluator_comments["P_External"])
                
                if p_delta_context:
                    with st.expander("P_Delta Context"):
                        st.json(p_delta_context)
                    with st.expander("P_Delta Section Title"):
                        st.text(evaluator_section_titles["P_Delta"])
                    with st.expander("Generated P_Delta Comments"):
                        st.text(generated_evaluator_comments["P_Delta"])
                
                
                session_id = "st-" + str(uuid.uuid4())
                time_taken = time.time() - start_time
                        
                db_utils.log_analyzer_custom_evaluator_data(
                    user_id=None,
                    user_name=None,
                    session_id=session_id,
                    proposal_pdf_name=proposal_pdf_name,
                    proposal_summary_text=proposal_summary_text,
                    proposal_text=proposal_text,
                    nature_of_document=nature,
                    organization_id=organization_id,
                    tor_pdf_name=tor_pdf_name,
                    tor_text=tor_text,
                    tor_summary_text=tor_summary_text,
                    generated_analyze_prompts=generated_analyze_prompts,
                    generated_analyze_comments=generated_analyze_comments,
                    analyze_context_used=analyze_context_used,
                    generated_evaluator_prompts=generated_evaluator_prompts,
                    generated_evaluator_comments=generated_evaluator_comments,
                    time_taken=time_taken,
                    tokens_counter=evaluator_tokens_counter
                )
    except Exception as e:
        logger.error(f"Error in generate_evaluator_comments: {e}")
        st.error(f"An error occurred in generate_evaluator_comments: {e}")


def generate_evaluator_comments(st, nature):

    used_evaluator_prompts = st.session_state['evaluator_prompts']
    evaluator_flows = st.session_state['evaluator_flows']
    generated_evaluator_prompts = st.session_state['generated_evaluator_prompts']
    proposal_pdf_url = st.session_state['proposal_pdf_url']
    proposal_summary_text = st.session_state['proposal_summary_text']
    proposal_text = st.session_state['proposal_text']
    tor_summary_text = st.session_state['tor_summary_text']
    evaluator_dependencies = st.session_state["evaluator_dependencies"]
    
    evaluator_dependencies = common_utils.filter_dependencies(['P_Internal'], evaluator_dependencies)
    evaluator_prompts_ordered = common_utils.topological_sort(evaluator_dependencies)
    
    
    with st.expander("Proposal S3 Url"):
        st.markdown(f'<a href="{proposal_pdf_url}" target="_blank">{proposal_pdf_url}</a>', unsafe_allow_html=True)
    
    with st.expander("Used Evaluator Prompts"):
        st.json(used_evaluator_prompts)

    with st.expander("Evaluator Flows"):
        st.json(evaluator_flows)
        
    with st.expander("Proposal Summary"):
        st.text(proposal_summary_text)
    
    with st.expander("TOR Summary"):
        st.text(tor_summary_text)
        
    with st.expander("Generated Evaluator Prompts"):
        st.json(generated_evaluator_prompts)
    
        
    model_options = {'gpt-4-omni': 'gpt-4', 'gpt-4-snapshot': 'gpt-4'}
    selected_display_name = st.selectbox('Evaluate Model', ('gpt-4-snapshot',))
    model_key = model_options[selected_display_name]
    model = possible_models[model_key]
    section_titles = {'P_Internal': ''}
    
    evaluate_submit = st.button("Evaluate")
    
    if evaluate_submit:
        start_time = time.time()
        
        with st.spinner("Generating Evaluate Comments..."):
            generated_evaluator_comments, p_internal_context, p_internal_flows, _, _ = gpt_utils.generate_evaluate_comments_multiflow(proposal_text,proposal_summary_text,nature, "",evaluator_flows,model,generated_evaluator_prompts,section_titles,evaluator_dependencies, evaluator_prompts_ordered)

            
            with st.expander("Full Document Text"):
                st.text(proposal_text)
                
            if p_internal_context:
                with st.expander("P_Internal Context"):
                    st.json(p_internal_context)
                    
            with st.expander("Generated P_Internal Comments"):
                st.text(generated_evaluator_comments["P_Internal"])
            
            if p_internal_flows:
                for key, value in p_internal_flows.items():
                    # Display F1 and F2 comments
                    for sub_key, sub_value in value["prompt_flow_results"].items():
                        with st.expander(f"{sub_key} Comments"):
                            comment = sub_value["comments"]
                            if(isinstance(comment,str)):
                                comment = {"comments": [{comment}]}
                            st.json(comment)

                    # Display F1 and F2 prompts
                    for sub_key, sub_value in value["prompt_flow_prompts"].items():
                        with st.expander(f"{sub_key} Prompts"):
                            st.json(sub_value)
                    
                    # Display the F1 and F2 contexts
                    for sub_key, sub_value in value["prompt_flow_contexts"].items():
                        with st.expander(f"{sub_key} Contexts"):
                            st.json(sub_value)
                with st.expander("P_Internal Flows"):
                    st.json(p_internal_flows)
            

            session_id = "st-"+str(uuid.uuid4())           
            time_taken = time.time() - start_time
                
            #db_utils.log_analyzer_custom_evaluator_data(user_id=None,user_name=None,session_id=session_id,proposal_pdf_name=proposal_pdf_name,proposal_summary_text=proposal_summary_text,proposal_text=proposal_text,nature_of_document=nature,identity=identity,organization_id=organization_id,org_guideline_id=org_guideline_id,tor_pdf_name=tor_pdf_name,tor_text=tor_text,tor_summary_text=tor_summary_text,generated_analyze_prompts=generated_analyze_prompts,generated_analyze_comments=generated_analyze_comments,analyze_context_used=analyze_context_used,generated_evaluator_prompts=generated_evaluator_prompts,generated_evaluator_comments=generated_evaluator_comments,time_taken=time_taken, input_token=input_tokens, output_token=output_tokens, total_token=total_tokens)


def generate_analyzer_prompts(st, summary_text, nature, selected_prompts_labels, selected_model, use_example=False, user_role=None):
    try:
        prompts_associated = {}
        with st.spinner("Generating Analyzer Prompts..."):
            logger.info("Fetching Analyzer prompts using multithreading...")
            selected_prompts, wisdom, prompts_section_title, dependency_graph = db_utils.get_custom_analyzer_gpt_prompts_multithreaded(prompt_labels, nature, None)
            logger.info("Analyzer prompts fetched.")
            
            for key, value in selected_prompts.items():
                base_prompt = value[1]
                customization_prompt = value[0]
            
                prompts_associated[key] = {
                    "base_prompt": base_prompt,
                    "customization_prompt": customization_prompt
                }
            logger.info("Generating analyzer prompts...")
            generated_analyze_prompts = common_utils.generate_prompts(selected_model, nature, summary_text, use_example, selected_prompts)
            logger.info("Analyzer prompts generated.")
            
            st.session_state['analyze_prompts'] = prompts_associated
            st.session_state['generated_analyze_prompts'] = generated_analyze_prompts
            st.session_state['wisdom'] = wisdom
            st.session_state['prompts_section_title'] = prompts_section_title
            st.session_state['dependency_graph'] = dependency_graph
    except Exception as e:
        logger.error(f"Error in generate_analyzer_prompts: {e}")
        st.error(f"An error occurred in generate_analyzer_prompts: {e}")


def evaluator_panel(st):
    try:
        proposal_content, proposal_pdf_name, tor_content, tor_pdf_name, nature, organization_id, org_guideline_id, selected_prompt_labels, selected_flows, selected_model, submit = evaluator_prompts_panel(st)

        if not proposal_content:
            st.info("Please upload a proposal PDF file.")
            return -1
        
        if (len(selected_prompt_labels) == 1) and selected_prompt_labels[0] == "P_Delta":
            st.info("P_Delta comments cannot be generated without P_Internal and P_External")
            return -1

        if (len(selected_prompt_labels) == 2) and "P_Delta" in selected_prompt_labels:
            st.info("P_Delta comments cannot be generated without both P_Internal and P_External")
            return -1

        if submit:
            start_time = time.time()
            evaluator_tokens_counter = common_utils.get_evaluator_tokens_counter()
            st.session_state['evaluator_tokens_counter'] = evaluator_tokens_counter
            
            logger.info("Generating evaluator prompts...")
            generate_evaluator_prompts(st, proposal_content, proposal_pdf_name, tor_content, tor_pdf_name, nature, organization_id, org_guideline_id, selected_prompt_labels, selected_flows, selected_model)
            logger.info(f"Generated evaluator prompts in {time.time() - start_time:.2f} seconds.")
            # if 'proposal_summary_text' in st.session_state and "P_External" in selected_prompt_labels:
            #     logger.info("Proposal Summary Found! Generating Analyzer Prompts...")
            #     summary_text = st.session_state['proposal_summary_text']
            #     generate_analyzer_prompts(st, summary_text, nature, prompt_labels, selected_model, organization_id)
            #     logger.info(f"Analyzer prompts generated for P_External in {time.time() - start_time:.2f} seconds.")
            # else:
            #     st.session_state["generated_analyze_prompts"] = ""
            #     st.session_state['wisdom'] = ""
            #     st.session_state['analyze_prompts'] = ""
            #     st.session_state['prompts_section_title'] = ""
            #     st.session_state['dependency_graph'] = ""

        if 'generated_evaluator_prompts' in st.session_state and 'evaluator_flows' in st.session_state:
            generate_evaluator_comments(st, nature)
            
    except Exception as e:
        logger.error(f"Error in evaluator_panel: {e}")
        traceback.print_exc()
        st.error(f"An error occurred in evaluator_panel: {e}")



